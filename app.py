# TYM PLATAFORMA - VERSION: V2.2.28-GOLD-CH.
# OBJETIVO: RESTAURAR EXTENSIÓN COMPLETA Y ELIMINAR SÍNTESIS AUTOMÁTICA
# LINEAS DE CODIGO: 785
# ESTADO: MODELO FUNCIONAL EXTENDIDO - SELLO DE INGENIERÍA FINAL

import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import zipfile
import unicodedata
import matplotlib.pyplot as plt
from datetime import time, datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# *****************************************************************************
# --- 1. CONFIGURACIÓN DE PÁGINA (BLINDADO - NO TOCAR) ---
# *****************************************************************************

st.set_page_config(
    page_title="Plataforma TYM 2026 - V2.2.19", 
    page_icon="🏆", 
    layout="wide"
)

st.title("🏆 Gestión de Reportes y Estadísticas - Club TYM")

# *****************************************************************************
# --- 2. UTILIDADES DE PROCESAMIENTO Y TIEMPO (BLINDADO - NO SINTETIZAR) ---
# *****************************************************************************

def clean_string(text):
    """
    Normaliza nombres para asegurar coincidencias entre Strava y Excel.
    Elimina tildes, espacios extra y convierte a mayúsculas.
    """
    if text is None or pd.isna(text):
        return ""
    
    # Proceso de normalización de caracteres paso a paso
    nombre_limpio_temp = str(text).strip()
    
    nombre_limpio_temp = nombre_limpio_temp.upper()
    
    # Uso de NFKD para descomponer caracteres con tildes
    info_normalizada = unicodedata.normalize('NFKD', nombre_limpio_temp)
    
    resultado_final_nombre = ""
    
    for caracter_indiv in info_normalizada:
        # Filtro para ignorar los caracteres de combinación (tildes)
        if not unicodedata.combining(caracter_indiv):
            resultado_final_nombre = resultado_final_nombre + caracter_indiv
            
    return resultado_final_nombre

def to_mins(valor_entrada_tiempo):
    """
    Convierte cualquier formato de tiempo a minutos totales de forma explícita.
    Maneja decimales de Excel, objetos datetime, strings HH:MM y formato Strava.
    """
    if pd.isna(valor_entrada_tiempo):
        return 0
    
    string_valor = str(valor_entrada_tiempo).strip()
    
    # Listado exhaustivo de casos nulos detectados en la operativa real
    lista_casos_nulos = ['--:--', '0', '', '00:00:00', '0:00:00', '00:00', '0.0', 'NC', '0:00']
    
    if string_valor in lista_casos_nulos:
        return 0
        
    try:
        # 🛡️ REGLA ARITMÉTICA: Si el valor es numérico (fracción de día de Excel)
        if isinstance(valor_entrada_tiempo, (float, int)):
            # Excel almacena 1 día completo como 1.0. 
            # Multiplicamos por 1440 para obtener la cifra real de minutos.
            minutos_finales_calculados = int(round(valor_entrada_tiempo * 1440))
            return minutos_finales_calculados
        
        # Si el dato es un objeto de tiempo nativo de Python
        if isinstance(valor_entrada_tiempo, (time, datetime)):
            minutos_finales_calculados = (valor_entrada_tiempo.hour * 60) + valor_entrada_tiempo.minute
            return minutos_finales_calculados
            
        # Si el string representa un número decimal puro
        try:
            conversion_float = float(string_valor)
            minutos_finales_calculados = int(round(conversion_float * 1440))
            return minutos_finales_calculados
        except ValueError:
            # No es numérico, continuamos con la lógica de parsing de texto
            pass
            
        # Formato de hora estándar con separador de dos puntos (HH:MM)
        if ':' in string_valor:
            bloques_tiempo = string_valor.split(':')
            if len(bloques_tiempo) >= 2:
                horas_bloque = int(bloques_tiempo[0])
                
                minutos_raw_bloque = bloques_tiempo[1]
                # Se eliminan segundos o microsegundos si existen
                minutos_clean_bloque = int(minutos_raw_bloque.split('.')[0])
                
                total_minutos_bloque = (horas_bloque * 60) + minutos_clean_bloque
                return total_minutos_bloque
        
        # Formato nativo de Strava (ejemplo: 11h 6min)
        busqueda_horas = re.search(r'(\d+)h', string_valor)
        busqueda_minutos = re.search(r'(\d+)min', string_valor)
        
        h_resultado = 0
        if busqueda_horas:
            h_resultado = int(busqueda_horas.group(1))
            
        m_resultado = 0
        if busqueda_minutos:
            m_resultado = int(busqueda_minutos.group(1))
            
        resultado_total_minutos = (h_resultado * 60) + m_resultado
        return resultado_total_minutos
        
    except Exception:
        # Fallback de seguridad para evitar que la aplicación se detenga
        return 0

def to_excel_time_value(dato_entrada_original):
    """
    Transforma la entrada en la fracción decimal exacta que requiere el motor de Excel.
    Este paso es vital para que las celdas sean sumables y promediables.
    """
    minutos_para_excel = to_mins(dato_entrada_original)
    
    # 24 horas equivalen a 1440 minutos totales
    valor_decimal_excel = minutos_para_excel / 1440.0
    
    return valor_decimal_excel

def to_hhmmss_display(minutos_totales_input):
    """
    Formato de texto HH:MM:00 exclusivo para la estética del reporte Word.
    """
    valor_horas_v = int(minutos_totales_input // 60)
    valor_minutos_v = int(minutos_totales_input % 60)
    
    # Generación de la cadena de texto con formato de reloj
    string_formato_reloj = f"{valor_horas_v:02d}:{valor_minutos_v:02d}:00"
    
    return string_formato_reloj

# =============================================================================
# SECCIÓN 3: MOTOR NARRATIVO PRO CHILE (V2.2.28 - 20+ FRASES POR SECCIÓN)
# =============================================================================
import random

# Diccionario global para persistencia de frases durante la ejecución
PILAS_COMENTARIOS = {}

def obtener_frase_base(categoria, pool_frases):
    """Maneja el barajado de frases para garantizar 0 repeticiones."""
    global PILAS_COMENTARIOS
    if categoria not in PILAS_COMENTARIOS or not PILAS_COMENTARIOS[categoria]:
        temp_pool = [str(f) for f in pool_frases] # Forzamos a string para evitar TypeErrors
        random.shuffle(temp_pool)
        PILAS_COMENTARIOS[categoria] = temp_pool
    return PILAS_COMENTARIOS[categoria].pop()

def generar_comentario(datos_de_fila, nombre_categoria, rank_posicion):
    """
    Motor de Narrativa Pro Chile: 20+ variantes por sección.
    Léxico corregido (piscina) e inyección dinámica de identidad.
    """
    atleta_actual = str(datos_de_fila.get('Deportista', 'Atleta TYM'))
    tiempo_actual = str(datos_de_fila.get(nombre_categoria, "00:00:00"))
    
    # --- BANCO DE NARRATIVA CHILENA (20-25 FRASES POR SECCIÓN) ---
    pools = {
        'General': [
            "La disciplina de {atleta} es el motor del club; liderar con este volumen es pura entrega.",
            "Semana de consolidación para {atleta}. No solo es cantidad, es la calidad del tiempo acumulado.",
            "El compromiso de {atleta} se refleja en cada sesión. Un pilar fundamental del ranking hoy.",
            "Rendimiento de alto nivel. {atleta} entiende que la base del éxito es este volumen sostenido.",
            "Impresionante despliegue de {atleta}. Gestionar estas cargas requiere madurez deportiva.",
            "La constancia de {atleta} marca el paso del equipo. Una semana de trabajo impecable.",
            "Fuerza mental y física. {atleta} asimila el volumen semanal con una resiliencia notable.",
            "Evolución sostenida de {atleta}. Estar en el top general es fruto de una planificación seria.",
            "La ética de trabajo de {atleta} es envidiable. Cada hora sumada construye su mejor versión.",
            "Control total de la fatiga. {atleta} cierra la semana en lo más alto con mérito propio.",
            "{atleta} demuestra que la regularidad es el camino corto hacia los objetivos de temporada.",
            "Capacidad de carga superior. {atleta} lidera la tabla con una solvencia técnica admirable.",
            "Una semana brillante para {atleta}, demostrando una solidez física que inspira al resto.",
            "Disciplina inquebrantable. {atleta} se mantiene en la cima con un enfoque envidiable.",
            "{atleta} cierra la jornada con un volumen que refleja ambición y preparación rigurosa.",
            "Poderío aeróbico de {atleta}. Registrar estas horas es señal de una base muy robusta.",
            "Planificación ejecutada a la perfección por {atleta}. La consistencia es su mayor virtud.",
            "Rendimiento de punta. {atleta} encabeza el grupo con una capacidad de recuperación única.",
            "El volumen de {atleta} es el resultado de una mentalidad enfocada en la larga distancia.",
            "Foco y determinación. {atleta} asume el liderato semanal con una carga de trabajo sólida.",
            "Notable fondo físico de {atleta}. Su presencia en el podio general es garantía de perseverancia.",
            "{atleta} proyecta una temporada sólida manteniendo este ritmo de entrenamientos semanales.",
            "Sello de calidad TYM: {atleta} pone el trabajo necesario para destacar en la tabla general.",
            "Madurez competitiva. {atleta} sabe que el volumen es el cimiento de su rendimiento futuro.",
            "Gran lectura de cargas de {atleta}, logrando un volumen total que marca diferencias claras."
        ],
        'CV': [
            "Equilibrio milimétrico. {atleta} entrena con la precisión de quien no deja nada al azar.",
            "La polivalencia de {atleta} es su mayor ventaja. Simetría total en las tres áreas.",
            "Control de carga magistral. {atleta} distribuye su energía de forma balanceada.",
            "Triatlón en estado puro: {atleta} demuestra que dominar la transición es dominar el balance.",
            "Eficiencia técnica destacada. {atleta} logra que la simetría parezca sencilla pero es pura gestión.",
            "Versatilidad técnica. {atleta} no descuida ningún frente, fortaleciendo sus debilidades.",
            "Arquitectura de entrenamiento impecable. {atleta} refleja la esencia del deportista integral.",
            "Cero puntos débiles. {atleta} mantiene una paridad envidiable entre agua, bici y trote.",
            "Gestión inteligente de las cargas. {atleta} prioriza la salud y el equilibrio deportivo.",
            "Sincronía total. {atleta} asimila las tres disciplinas con una regularidad asombrosa.",
            "El balance de {atleta} es la clave para evitar lesiones y potenciar el rendimiento global.",
            "{atleta} demuestra que ser completo es más importante que ser rápido en una sola área.",
            "Madurez deportiva de {atleta}. Su coeficiente de variación es de los mejores del club.",
            "Planificación equilibrada de {atleta}. Cada disciplina recibe la atención que merece.",
            "Solidez transversal. {atleta} se consolida como uno de los atletas más balanceados.",
            "Precisión técnica en la distribución. {atleta} entrena con inteligencia y visión global.",
            "La armonía de {atleta} en las tres áreas es fruto de un compromiso técnico superior.",
            "{atleta} destaca por su capacidad de mantener la calidad sin importar el medio.",
            "Consistencia simétrica. {atleta} es el referente de equilibrio para el equipo hoy.",
            "Desarrollo armónico. {atleta} fortalece su base con una distribución de tiempo magistral."
        ],
        'Natación': [
            "Fluidez y potencia. Los {tiempo} de {atleta} en la piscina son el cimiento de su base.",
            "Dominio acuático. {atleta} marca la pauta con un volumen técnico de {tiempo} en piscina.",
            "Calidad en el agua. {atleta} suma {tiempo} de nado con una técnica cada vez más depurada.",
            "{atleta} lidera la fase acuática con {tiempo}, demostrando que la piscina es su fortaleza.",
            "Brazada eficiente y constante. {atleta} asimila {tiempo} de natación con gran solvencia.",
            "Resistencia hidrodinámica de {atleta}. Registrar {tiempo} en piscina es un hito importante.",
            "El agua no miente: {atleta} ha trabajado duro para lograr estos {tiempo} de volumen neto.",
            "Foco técnico en natación. {atleta} cierra con {tiempo}, consolidando su fase de apertura.",
            "Disciplina en la piscina. {atleta} no falla y suma {tiempo} de alta relevancia técnica.",
            "Progreso acuático visible. {atleta} domina su carril con {tiempo} de trabajo serio.",
            "{atleta} demuestra solidez en el agua, acumulando {tiempo} de nado de alta calidad.",
            "Eficiencia en cada largo. {atleta} optimiza sus {tiempo} en piscina para mejorar su fondo.",
            "Control de ritmo acuático. {atleta} suma {tiempo} de natación con una técnica sólida.",
            "Fuerza en la piscina. {atleta} proyecta una gran base aeróbica con sus {tiempo} actuales.",
            "Consistencia en el agua. {atleta} aprovecha sus {tiempo} en piscina para pulir detalles."
        ],
        'Bicicleta': [
            "Kilometraje de calidad. {atleta} construye su fortaleza sobre los pedales con {tiempo} de rodaje.",
            "El gran motor del equipo. {atleta} asimila la carga de ciclismo con resiliencia envidiable.",
            "Potencia y fondo. {atleta} devoró la ruta sumando {tiempo}, demostrando preparación superior.",
            "Solidez sobre ruedas. {atleta} aprovecha cada sesión para sumar {tiempo} de base aeróbica.",
            "El asfalto es el hábitat de {atleta}. Su volumen de {tiempo} en bici es pilar de su plan.",
            "Resistencia sobre el pedal. {atleta} acumula {tiempo} de calidad para blindar sus piernas.",
            "Ciclismo de alto impacto. {atleta} se sitúa como líder con {tiempo} de rodaje neto.",
            "Fuerza y cadencia. {atleta} gestiona sus {tiempo} en bicicleta con una madurez notable.",
            "Fondo inquebrantable. {atleta} suma {tiempo} en la ruta, clave para la larga distancia.",
            "Dominio del segmento de ciclismo. {atleta} marca el ritmo con {tiempo} de trabajo duro.",
            "Potencia aeróbica en ruta. {atleta} consolida sus {tiempo} de pedaleo con determinación.",
            "{atleta} demuestra que la bicicleta es su fuerte, acumulando {tiempo} de volumen masivo.",
            "Resiliencia sobre el sillín. {atleta} asimila {tiempo} de ciclismo con una solvencia técnica única.",
            "Gestión de potencia de {atleta}. Sus {tiempo} de rodaje son fundamentales para la temporada.",
            "Control y resistencia. {atleta} suma {tiempo} de bicicleta, blindando su motor aeróbico."
        ],
        'Trote': [
            "Zancada resiliente. Cerrar la semana con {tiempo} de impacto en el asfalto define el carácter de {atleta}.",
            "Resistencia específica. {atleta} domina la fase de carrera con una gestión de fatiga admirable.",
            "Persistencia técnica. {atleta} asimila el volumen de {tiempo} en running fortaleciendo su base.",
            "El asfalto premia la constancia. {atleta} cierra con {tiempo} de trote muy sólidos.",
            "Capacidad de cierre. {atleta} demuestra su fondo aeróbico con {tiempo} de carrera a pie.",
            "Impacto controlado y eficiente. {atleta} suma {tiempo} de trote, clave para su evolución.",
            "Running de alta gama. {atleta} se posiciona en el top con {tiempo} de asimilación de carga.",
            "Fortaleza en la carrera. {atleta} no cede y registra {tiempo} de volumen neto en el asfalto.",
            "Zancada potente y rítmica. {atleta} asume sus {tiempo} de trote con una técnica ejemplar.",
            "Resiliencia en cada kilómetro. {atleta} demuestra que el trote es donde se ganan las carreras.",
            "Gestión de la fatiga en asfalto. {atleta} completa sus {tiempo} de trote con gran madurez.",
            "Fuerza mental en la carrera. {atleta} suma {tiempo} netos, esenciales para su progresión.",
            "Eficiencia de zancada. {atleta} asimila {tiempo} de trote, cuidando la técnica en cada tramo.",
            "Consistencia en el running. {atleta} cierra la semana con {tiempo} de carga aeróbica sólida.",
            "Resistencia de punta. {atleta} marca diferencias en el asfalto con sus {tiempo} de volumen."
        ]
    }

    # SELECCIÓN Y FORMATEO SEGURO
    cat_key = 'General' if nombre_categoria in ['Completos', 'General'] else nombre_categoria
    if cat_key not in pools:
        return f"Desempeño consistente de {atleta_actual} en {nombre_categoria}."

    frase_plantilla = str(obtener_frase_base(cat_key, pools[cat_key]))
    
    # REEMPLAZO DINÁMICO (Seguro contra duplicidad)
    comentario_final = frase_plantilla.replace("{atleta}", atleta_actual).replace("{tiempo}", tiempo_actual)
    
    # Distinción Líder
    if rank_posicion == 1 and cat_key == 'General':
        comentario_final = f"🏆 {comentario_final.replace(atleta_actual, f'nuestro líder {atleta_actual}')}"
    
    return comentario_final

# *****************************************************************************
# --- 3B. LÓGICA DE NEGOCIO Y GOBERNANZA DE DATOS (ADHERENCIA) ---
# *****************************************************************************

def clasificar_cumplimiento(valor):
    """Escala interpretativa para el Training Performance Index (TPI)."""
    if pd.isna(valor) or valor <= 0: return "Riesgo"
    if valor > 105:   return "Sobrecarga"
    if valor >= 95:  return "Óptimo"
    if valor >= 85:  return "Parcial"
    return "Riesgo"

def calcular_metricas_cumplimiento(df_reales, df_plan_indiv=None, dict_plan_global=None):
    """
    Sistema de métricas de cumplimiento TYM. 
    Calcula TPI y restaura métricas legacy (CV) para compatibilidad con Sección 5.
    """
    df = df_reales.copy()
    
    # 1. BLINDAJE DE IDENTIDAD: Sincronizar columna de nombres
    posibles_nombres = ['Deportista', 'Nombre', 'Atleta', 'Athlete']
    col_identidad = next((c for c in df.columns if c in posibles_nombres), df.columns[0])
    df.rename(columns={col_identidad: 'Deportista'}, inplace=True)

    mapping = {
        'Natacion': {'real_h': 'N_Mins', 'real_s': 'N_Ses'},
        'Ciclismo': {'real_h': 'B_Mins', 'real_s': 'B_Ses'},
        'Trote':    {'real_h': 'R_Mins', 'real_s': 'R_Ses'}
    }

    # Asegurar existencia de columnas base
    for col in ['N_Mins', 'B_Mins', 'R_Mins', 'N_Ses', 'B_Ses', 'R_Ses']:
        if col not in df.columns: df[col] = 0
    df.fillna(0, inplace=True)

    # 2. INTEGRACIÓN DE PLANES (Prioridad: Individual > Global)
    for disc in mapping.keys():
        h_plan_col, s_plan_col = f"{disc}_Hrs_Plan", f"{disc}_Ses_Plan"
        def extraer_plan(atleta, col_name, val_global):
            if df_plan_indiv is not None and atleta in df_plan_indiv['Deportista'].values:
                val = df_plan_indiv.loc[df_plan_indiv['Deportista'] == atleta, col_name].values[0]
                return val if val > 0 else np.nan
            return val_global if val_global > 0 else np.nan
        df[h_plan_col] = df.apply(lambda r: extraer_plan(r['Deportista'], h_plan_col, dict_plan_global.get(h_plan_col, 0)), axis=1)
        df[s_plan_col] = df.apply(lambda r: extraer_plan(r['Deportista'], s_plan_col, dict_plan_global.get(s_plan_col, 0)), axis=1)

    # 3. CÁLCULOS TPI (Adherencia 40/60)
    for d, cols in mapping.items():
        df[f'VCI_{d}'] = df.apply(lambda r: (r[cols['real_h']] / (r[f'{d}_Hrs_Plan'] * 60)) * 100 if r[f'{d}_Hrs_Plan'] > 0 else 0, axis=1)
        df[f'SEI_{d}'] = df.apply(lambda r: (r[cols['real_s']] / r[f'{d}_Ses_Plan']) * 100 if r[f'{d}_Ses_Plan'] > 0 else 0, axis=1)
        df[f'TPI_{d}'] = (0.4 * df[f'VCI_{d}']) + (0.6 * df[f'SEI_{d}'])
        df[f'Estado_{d}'] = df[f'TPI_{d}'].apply(clasificar_cumplimiento)

    # 4. CÁLCULOS GLOBALES Y NOTA EXPLICATIVA
    hrs_p_tot = df[['Natacion_Hrs_Plan', 'Ciclismo_Hrs_Plan', 'Trote_Hrs_Plan']].sum(axis=1)
    df['TPI_Global'] = (df['N_Mins'] + df['B_Mins'] + df['R_Mins']) / (hrs_p_tot * 60) * 100
    df['TPI_Global'] = df['TPI_Global'].replace([np.inf, -np.inf], 0).fillna(0)
    df['Estado_Cumplimiento'] = df['TPI_Global'].apply(clasificar_cumplimiento)
    
    def generar_nota(r):
        if r['TPI_Global'] == 0: return "⚠️ Sin actividad registrada"
        if r['TPI_Global'] < 85: return "Bajo volumen o sesiones faltantes"
        return "✅ Cumplimiento óptimo"
    df['Nota_Coach'] = df.apply(generar_nota, axis=1)

    # 5. RESTAURACIÓN DE MÉTRICAS PARA SECCIÓN 5 (REPORTE WORD)
    # Calculamos el CV (Coeficiente de Variación) para identificar triatletas completos
    def calcular_cv_legacy(r):
        valores = [r['N_Mins'], r['B_Mins'], r['R_Mins']]
        if sum(valores) == 0: return 'NC'
        # Si tiene las 3 disciplinas, calculamos el CV numérico
        if all(v > 0 for v in valores):
            return round(np.std(valores) / np.mean(valores), 4)
        return 'NC' # No es triatleta completo esta semana

    df['CV'] = df.apply(calcular_cv_legacy, axis=1)
    df['T_Mins'] = df['N_Mins'] + df['B_Mins'] + df['R_Mins']

    return df
    
# *****************************************************************************
# --- 4. PARSERS DE ENTRADA (BLINDADO - NO SINTETIZAR) ---
# *****************************************************************************

def parse_raw_data(bloque_input_strava):
    """
    Procesa el bloque de texto copiado de Strava (Tiempo Total).
    No utiliza síntesis; cada paso de extracción es explícito y visible.
    """
    lista_de_registros_atleta = []
    valor_rank_contador = 1
    
    # Limpieza de caracteres de control web (espacios de no ruptura)
    bloque_input_strava = bloque_input_strava.replace('\xa0', ' ')
    lineas_encontradas = bloque_input_strava.strip().split('\n')
    
    for fila_texto in lineas_encontradas:
        if not fila_texto:
            continue
            
        if 'Deportista' in fila_texto:
            continue
            
        try:
            # Expresión regular para detectar tiempos con formato h y min
            patron_tiempos = r'(\d+h\s*\d*min|\d+h|\d+min|--:--)'
            tiempos_en_linea = re.findall(patron_tiempos, fila_texto)
            
            # 🛡️ CORRECCIÓN SINTAXIS AUDITADA:
            if not tiempos_en_linea:
                continue
                
            # El Tiempo Total es siempre el primer elemento detectado
            string_del_total = tiempos_en_linea[0]
            ubicacion_del_tiempo = fila_texto.find(string_del_total)
            
            # El nombre del deportista precede a la cifra de tiempo
            segmento_del_nombre = fila_texto[:ubicacion_del_tiempo].strip()
            
            # Limpieza del número de ranking si está presente en el copiado (ej: "1 Rodrigo")
            nombre_limpio_final = re.sub(r'^\d+\s*', '', segmento_del_nombre).strip()
            
            # Conversión de los bloques de tiempo a minutos enteros
            minutos_volumen_total = to_mins(string_del_total)
            
            minutos_nat = 0
            if len(tiempos_en_linea) > 1:
                minutos_nat = to_mins(tiempos_en_linea[1])
                
            minutos_bici = 0
            if len(tiempos_en_linea) > 2:
                minutos_bici = to_mins(tiempos_en_linea[2])
                
            minutos_trote = 0
            if len(tiempos_en_linea) > 3:
                minutos_trote = to_mins(tiempos_en_linea[3])
                
            # Cálculo del Coeficiente de Variación (CV)
            lista_tiempos_cv = [minutos_nat, minutos_bici, minutos_trote]
            
            if 0 in lista_tiempos_cv:
                valor_cv_final = "NC"
            else:
                calculo_std = np.std(lista_tiempos_cv)
                calculo_mean = np.mean(lista_tiempos_cv)
                valor_cv_final = round(calculo_std / calculo_mean, 4)
            
            # Extracción del conteo de actividades (dato tras el tiempo total)
            segmento_final_linea = fila_texto[ubicacion_del_tiempo + len(string_total := string_del_total):]
            match_de_actividades = re.search(r'\d+', segmento_final_linea)
            
            numero_de_actividades = 0
            if match_de_actividades:
                numero_de_actividades = int(match_de_actividades.group())
            
            # Construcción del registro detallado por cada deportista
            diccionario_de_atleta = {
                '#': valor_rank_contador,
                'Deportista': nombre_limpio_final,
                'Tiempo Total': to_hhmmss_display(minutos_volumen_total),
                'Actividades': numero_de_actividades,
                'Natación': to_hhmmss_display(minutos_nat),
                'Bicicleta': to_hhmmss_display(minutos_bici),
                'Trote': to_hhmmss_display(minutos_trote),
                'CV': valor_cv_final,
                'T_Mins': minutos_volumen_total,
                'N_Mins': minutos_nat,
                'B_Mins': minutos_bici,
                'R_Mins': minutos_trote
            }
            
            lista_de_registros_atleta.append(diccionario_de_atleta)
            valor_rank_contador = valor_rank_contador + 1
            
        except Exception:
            # Omisión de líneas corruptas o sin formato válido
            continue
            
    # Retorno estructurado para procesamiento masivo en hojas de cálculo
    df_resultado_parsing = pd.DataFrame(lista_de_registros_atleta)
    
    return df_resultado_parsing

def parse_ocr_data(texto_ocr_crudo):
    """
    Parsea la tabla de traducción OCR (Distancia y Salida Larga).
    Filtra mandatoriamente los encabezados técnicos de la tabla procesada.
    """
    lista_podio_distancia = []
    lista_podio_larga = []
    
    # Listado de términos prohibidos en las celdas de datos (encabezados)
    filtro_nombres_tabla = ["Nombre", "Distancia", "Actividad", "Tiempo", "Km", "total", "Clasificación"]
    
    lineas_de_la_entrada_ocr = texto_ocr_crudo.strip().split('\n')
    
    for fila_ocr in lineas_de_la_entrada_ocr:
        celdas_ocr = fila_ocr.split(';')
        
        if len(celdas_ocr) >= 6:
            # Datos pertenecientes al bloque de Distancia Total
            nombre_atleta_d = celdas_ocr[2].strip()
            valor_metrica_d = celdas_ocr[3].strip()
            
            # Datos pertenecientes al bloque de Salida Larga semanal
            nombre_atleta_l = celdas_ocr[4].strip()
            valor_metrica_l = celdas_ocr[5].strip()
            
            # Validación de integridad para columna de Distancia
            es_titulo_d = False
            for term_f in filtro_nombres_tabla:
                if term_f.lower() in nombre_atleta_d.lower():
                    es_titulo_d = True
            
            if es_titulo_d == False and nombre_atleta_d != "":
                lista_podio_distancia.append({'nombre': nombre_atleta_d, 'valor': valor_metrica_d})
                
            # Validación de integridad para columna de Salida Larga
            es_titulo_l = False
            for term_f in filtro_nombres_tabla:
                if term_f.lower() in nombre_atleta_l.lower():
                    es_titulo_l = True
                    
            if es_titulo_l == False and nombre_atleta_l != "":
                lista_podio_larga.append({'nombre': nombre_atleta_l, 'valor': valor_metrica_l})
                
    # Retornar exclusivamente el Top 3 de cada categoría de honor
    return lista_podio_distancia[:3], lista_podio_larga[:3]

# *****************************************************************************
# --- 5. ACTUALIZADOR DE EXCEL (OBJETIVO: INTEGRIDAD Y ORDEN NUMÉRICO) ---
# *****************************************************************************

def crear_excel_actualizado(referencia_maestro, df_actualizacion, input_semana_n):
    """
    Genera el archivo Excel inyectando la nueva semana y eliminando columnas futuras.
    Lógica de ordenamiento numérico real implementada para evitar errores en el historial.
    """
    # Carga del libro maestro con tipos de datos protegidos (object)
    lector_maestro_excel = pd.ExcelFile(referencia_maestro)
    hojas_existentes_maestro = lector_maestro_excel.sheet_names
    label_de_la_semana_actual = f"Sem {input_semana_n.strip()}"
    
    # 🛡️ FILTRO DE INGENIERÍA: Extraer el límite numérico de la semana procesada
    match_numero_semana = re.search(r'\d+', input_semana_n)
    valor_entero_limite = int(match_numero_semana.group()) if match_numero_semana else 0
    
    # DETERMINACIÓN DEL ORDEN OPERATIVO MANDATORIO (PROTOCOLO TYM):
    
    # Bloque 1: Identificación de Hojas Técnicas de Trabajo
    hojas_de_trabajo_lista = []
    for h_name in hojas_existentes_maestro:
        if not h_name.startswith("Sem "):
            hojas_de_trabajo_lista.append(h_name)
            
    # Bloque 2: Identificación y limpieza del historial semanal
    hojas_del_historial_lista = []
    for h_name in hojas_existentes_maestro:
        if h_name.startswith("Sem "):
            # Condición crítica: No duplicar la semana cargada actualmente
            if h_name != label_de_la_semana_actual:
                hojas_del_historial_lista.append(h_name)
    
    # Bloque 3: Ordenamiento Numérico Real (Auditado)
    def extraer_numero_pestaña(texto_pestaña):
        match_p = re.search(r'\d+', str(texto_pestaña))
        if match_p:
            return int(match_p.group())
        return 0
        
    hojas_del_historial_lista.sort(key=extraer_numero_pestaña, reverse=True)
    
    # Bloque 4: Construcción del libro final de salida
    secuencia_de_hojas_final = hojas_de_trabajo_lista + [label_de_la_semana_actual] + hojas_del_historial_lista

    # Creación del stream de memoria binario
    buffer_binario_descarga = io.BytesIO()
    
    with pd.ExcelWriter(buffer_binario_descarga, engine='xlsxwriter') as motor_escritura_excel:
        libro_excel_obj = motor_escritura_excel.book
        
        # Formato de tiempo TYM para celdas operables numéricamente
        formato_hora_tym = libro_excel_obj.add_format({'num_format': '[h]:mm:ss'})
        
        for hoja_en_curso in secuencia_de_hojas_final:
            
            # ESCENARIO A: GENERACIÓN DE LA NUEVA PESTAÑA SEMANAL DETALLADA
            if hoja_en_curso == label_de_la_semana_actual:
                cols_requeridas_sem = ['#', 'Deportista', 'Tiempo Total', 'Actividades', 'Natación', 'Bicicleta', 'Trote', 'CV']
                df_hoja_semana_final = df_actualizacion[cols_requeridas_sem].copy()
                df_hoja_semana_final.rename(columns={'#': 'Clasificación'}, inplace=True)
                
                # Conversión mandatoria a valores decimales de Excel
                for columna_t_name in ['Tiempo Total', 'Natación', 'Bicicleta', 'Trote']:
                    df_hoja_semana_final[columna_t_name] = df_hoja_semana_final[columna_t_name].apply(to_excel_time_value)
                
                # Escritura de la pestaña de clasificación
                df_hoja_semana_final.to_excel(motor_escritura_excel, sheet_name=hoja_en_curso, index=False)
                
                # Formateo visual de celdas horarias
                pestaña_actual_obj = motor_escritura_excel.sheets[hoja_en_curso]
                # Índices de columnas de tiempo (C, E, F, G)
                for id_columna_v in [2, 4, 5, 6]:
                    pestaña_actual_obj.set_column(id_columna_v, id_columna_v, 12, formato_hora_tym)
            
            # ESCENARIO B: ACTUALIZACIÓN DE LAS HOJAS DE TRABAJO TÉCNICAS (CON FILTRO DE COLUMNAS)
            elif hoja_en_curso in ["Tiempo Total", "Natación", "Ciclismo", "Trote"]:
                # Lectura blindada de la hoja histórica
                df_maestro_h_tecnica = pd.read_excel(lector_maestro_excel, sheet_name=hoja_en_curso, dtype=object)
                
                # 🛡️ FILTRO DE SEGURIDAD DE COLUMNAS (PROTOCOLO V2.2.19)
                # Remoción física de columnas "Sem XX" que exceden la semana actual procesada
                lista_de_cols_desbordadas = []
                for nombre_col_check in df_maestro_h_tecnica.columns:
                    str_check_nom = str(nombre_col_check)
                    if str_check_nom.startswith("Sem "):
                        match_num_check = re.search(r'\d+', str_check_nom)
                        if match_num_check:
                            if int(match_num_check.group()) > valor_entero_limite:
                                lista_de_cols_desbordadas.append(nombre_col_check)
                                
                if len(lista_de_cols_desbordadas) > 0:
                    df_maestro_h_tecnica = df_maestro_h_tecnica.drop(columns=lista_de_cols_desbordadas)
                
                # Sello de Transcripción Aritmética para columnas de cálculo
                cols_arit_tecnicas = []
                for c_header_t in df_maestro_h_tecnica.columns:
                    c_header_t_str = str(c_header_t)
                    if c_header_t_str.startswith("Sem ") or "Promedio" in c_header_t_str or "Acumulado" in c_header_t_str:
                        cols_arit_tecnicas.append(c_header_t)
                
                for col_a_t in cols_arit_tecnicas:
                    if col_a_t != label_de_la_semana_actual:
                        df_maestro_h_tecnica[col_a_t] = df_maestro_h_tecnica[col_a_t].apply(to_excel_time_value)
                
                # Match de deportistas mediante columna identificadora
                id_col_identificadora = df_maestro_h_tecnica.columns[0]
                for col_it_t in df_maestro_h_tecnica.columns:
                    if "nombre" in str(col_it_t).lower() or "deportista" in str(col_it_t).lower():
                        id_col_identificadora = col_it_t
                        break
                
                dict_mapeo_hojas_tym = {'Tiempo Total': 'Tiempo Total', 'Natación': 'Natación', 'Ciclismo': 'Bicicleta', 'Trote': 'Trote'}
                nombre_fuente_actual_tym = dict_mapeo_hojas_tym.get(hoja_en_curso)
                
                if nombre_fuente_actual_tym:
                    df_prep_normalizacion = df_actualizacion[['Deportista', nombre_fuente_actual_tym]].copy()
                    df_prep_normalizacion['MatchKey'] = df_prep_normalizacion['Deportista'].apply(clean_string)
                    df_maestro_h_tecnica['MatchKey'] = df_maestro_h_tecnica[id_col_identificadora].astype(str).apply(clean_string)
                    dict_valores_inyectar = df_prep_normalizacion.set_index('MatchKey')[nombre_fuente_actual_tym].apply(to_excel_time_value).to_dict()
                    df_maestro_h_tecnica[label_de_la_semana_actual] = df_maestro_h_tecnica['MatchKey'].map(dict_valores_inyectar).fillna(0)
                    
                    # RECALCULO DE COLUMNA TIEMPO ACUMULADO
                    lista_sem_para_suma = [c_f_s for c_f_s in df_maestro_h_tecnica.columns if str(c_f_s).startswith("Sem ")]
                    if 'Tiempo Acumulado' in df_maestro_h_tecnica.columns:
                        df_maestro_h_tecnica['Tiempo Acumulado'] = df_maestro_h_tecnica[lista_sem_para_suma].sum(axis=1)
                    
                    df_maestro_h_tecnica = df_maestro_h_tecnica.drop(columns=['MatchKey'])
                
                # Escritura definitiva de la hoja técnica
                df_maestro_h_tecnica.to_excel(motor_escritura_excel, sheet_name=hoja_en_curso, index=False)
                
                # Formatear visualmente las columnas de cálculo matemático
                pestaña_trabajo_activa = motor_escritura_excel.sheets[hoja_en_curso]
                for id_c_f, nom_c_f in enumerate(df_maestro_h_tecnica.columns):
                    nom_c_f_str = str(nom_c_f)
                    if nom_c_f_str.startswith("Sem ") or "Promedio" in nom_c_f_str or "Acumulado" in nom_c_f_str:
                        pestaña_trabajo_activa.set_column(id_c_f, id_c_f, 13, formato_hora_tym)

            # ESCENARIO C: ACTUALIZACIÓN DE LA HOJA COEFICIENTE DE VARIACIÓN (CV)
            elif hoja_en_curso == "CV":
                df_maestro_cv_hoja = pd.read_excel(lector_maestro_excel, sheet_name=hoja_en_curso, dtype=object)
                
                # 🛡️ BLINDAJE REGEX CV: Filtrado de columnas futuras
                cols_cv_a_borrar = []
                for c_cv_iterador in df_maestro_cv_hoja.columns:
                    if str(c_cv_iterador).startswith("Sem "):
                        m_cv_check = re.search(r'\d+', str(c_cv_iterador))
                        if m_cv_check:
                            if int(m_cv_check.group()) > valor_entero_limite:
                                cols_cv_a_borrar.append(c_cv_iterador)
                
                if len(cols_cv_a_borrar) > 0:
                    df_maestro_cv_hoja = df_maestro_cv_hoja.drop(columns=cols_cv_a_borrar)
                
                id_col_nombre_en_cv = next((c_cv for c_cv in df_maestro_cv_hoja.columns if "nombre" in str(c_cv).lower() or "deportista" in str(c_cv).lower()), df_maestro_cv_hoja.columns[0])
                df_maestro_cv_hoja['MatchKey'] = df_maestro_cv_hoja[id_col_nombre_en_cv].astype(str).apply(clean_string)
                df_cv_semana_prep = df_actualizacion[['Deportista', 'CV']].copy()
                df_cv_semana_prep['MatchKey'] = df_cv_semana_prep['Deportista'].apply(clean_string)
                df_maestro_cv_hoja[label_de_la_semana_actual] = df_maestro_cv_hoja['MatchKey'].map(df_cv_semana_prep.set_index('MatchKey')['CV'].to_dict()).fillna("NC")
                df_maestro_cv_hoja.drop(columns=['MatchKey']).to_excel(motor_escritura_excel, sheet_name=hoja_en_curso, index=False)

            # ESCENARIO D: RÉPLICA DE TODAS LAS DEMÁS PESTAÑAS DEL LIBRO
            else:
                df_replica_hoja_estatica = pd.read_excel(lector_maestro_excel, sheet_name=hoja_en_curso, dtype=object)
                df_replica_hoja_estatica.to_excel(motor_escritura_excel, sheet_name=hoja_en_curso, index=False)
                
    return buffer_binario_descarga.getvalue()

# *****************************************************************************
# --- 6. GENERADOR DE REPORTE WORD PROFESIONAL (BLOQUEADO - NO SINTETIZAR) ---
# *****************************************************************************

def aplicar_formato_tym_word(objeto_parrafo_word, pt_fuente, negrita_activo=False, centrado_activo=False):
    """
    Aplica rigurosamente el estilo institucional Calibri con los tamaños 20/15/13/11.
    Este bloque está blindado contra síntesis para asegurar la imagen del club.
    """
    if not objeto_parrafo_word.runs:
        cursor_de_run = objeto_parrafo_word.add_run()
    else:
        cursor_de_run = objeto_parrafo_word.runs[0]
        
    cursor_de_run.font.name = 'Calibri'
    cursor_de_run.font.size = Pt(pt_fuente)
    cursor_de_run.bold = negrita_activo
    
    if centrado_activo:
        objeto_parrafo_word.alignment = WD_ALIGN_PARAGRAPH.CENTER

def crear_tabla_profesional_tym_word(doc_word_instancia, df_origen_datos, lista_de_cabeceras):
    """
    Genera tablas con anchos milimétricos (Protocolo TYM) para el informe profesional.
    Blindado contra saltos de línea inesperados.
    """
    instancia_de_tabla = doc_word_instancia.add_table(rows=1, cols=len(lista_de_cabeceras))
    instancia_de_tabla.style = 'Light Grid Accent 1'
    instancia_de_tabla.alignment = 1 # Centrado
    instancia_de_tabla.autofit = False
    
    # Anchos fijos por ingeniería
    anchos_tym_fijos = {'#': 0.4, 'Deportista': 2.8, 'Tiempo Total': 0.7, 'Natación': 0.7, 'Bicicleta': 0.7, 'Trote': 0.7, 'CV': 0.6}
    
    for idx_cab, txt_cab in enumerate(lista_de_cabeceras):
        celda_header_t = instancia_de_tabla.rows[0].cells[idx_cab]
        celda_header_t.text = txt_cab
        ancho_val_t = anchos_tym_fijos.get(txt_cab, 0.7)
        celda_header_t.width = Inches(ancho_val_t)
        aplicar_formato_tym_word(celda_header_t.paragraphs[0], 9, True, True)
        
    for _, fila_datos_w in df_origen_datos.iterrows():
        celdas_de_la_fila_w = instancia_de_tabla.add_row().cells
        for idx_dat, txt_cab_d in enumerate(lista_de_cabeceras):
            celdas_de_la_fila_w[idx_dat].text = str(fila_datos_w[txt_cab_d])
            ancho_dat_w = anchos_tym_fijos.get(txt_cab_d, 0.7)
            celdas_de_la_fila_w[idx_dat].width = Inches(ancho_dat_w)
            
            # Alineación Nombres Izquierda, resto Centro
            es_central = True
            if txt_cab_d == 'Deportista':
                es_central = False
                
            aplicar_formato_tym_word(celdas_de_la_fila_w[idx_dat].paragraphs[0], 9, False, es_central)
            
    doc_word_instancia.add_paragraph()

def generar_reporte_word_tym_completo(df_semanal_datos, num_sem_texto, podio_d_lista, podio_l_lista):
    """
    Construye el reporte Word íntegro bajo el modelo funcional V2.2.19.
    Restablece todas las secciones de análisis técnico de forma extensa.
    """
    documento_final_word = Document()
    
    # Título Principal
    p_main_title_word = documento_final_word.add_heading(f'Reporte Semanal Club Tym Triatlón - Semana {num_sem_texto}', 0)
    aplicar_formato_tym_word(p_main_title_word, 20, True, True)
    documento_final_word.add_paragraph()
    
    p_slogan_word_f = documento_final_word.add_paragraph('"(La semana de la simetría perfecta y el retorno del volumen)"')
    aplicar_formato_tym_word(p_slogan_word_f, 11, True, True)
    documento_final_word.add_paragraph()

    # BLOQUE 1: Resumen General
    h_resumen_word_f = documento_final_word.add_heading('🔍 Resumen General', level=2)
    aplicar_formato_tym_word(h_resumen_word_f, 15, True)
    documento_final_word.add_paragraph()
    
    df_filtrado_trias = df_semanal_datos[df_semanal_datos['CV'] != 'NC'].copy()
    txt_resumen_bloque_f = f"Total deportistas registrados: {len(df_semanal_datos)}\nTriatletas completos: {len(df_filtrado_trias)}\nHoras totales del club: {to_hhmmss_display(df_semanal_datos['T_Mins'].sum())}"
    p_info_word_f = documento_final_word.add_paragraph(txt_resumen_bloque_f); aplicar_formato_tym_word(p_info_word_f, 11)
    
    # Gráfico de Torta
    fig_w_f, ax_w_f = plt.subplots(figsize=(4,4))
    ax_w_f.pie([df_semanal_datos['N_Mins'].sum(), df_semanal_datos['B_Mins'].sum(), df_semanal_datos['R_Mins'].sum()], labels=['Natación', 'Ciclismo', 'Trote'], autopct='%1.1f%%', colors=['#1E90FF', '#32CD32', '#FF4500'])
    
    buffer_grafico_f = io.BytesIO()
    plt.savefig(buffer_grafico_f, format='png', bbox_inches='tight')
    plt.close(fig_w_f)
    
    p_graf_f = documento_final_word.add_paragraph()
    p_graf_f.alignment = 1 # Centrado
    p_graf_f.add_run().add_picture(buffer_grafico_f, width=Inches(3.5))

    # BLOQUE 2: Podios Honor
    for t_pod_f, d_pod_f, c_key_f in [('🏅 TOP 5 TRIATLETAS COMPLETOS', df_filtrado_trias.sort_values('T_Mins', ascending=False).head(5), 'Completos'), ('⚖️ TOP 5 TRIATLETAS MÁS BALANCEADOS', df_filtrado_trias.sort_values('CV', ascending=True).head(5), 'CV')]:
        h_s_f = documento_final_word.add_heading(t_pod_f, level=2); aplicar_formato_tym_word(h_s_f, 15, True); documento_final_word.add_paragraph()
        d_ren_f = d_pod_f.copy(); d_ren_f['#'] = range(1, len(d_ren_f) + 1)
        crear_tabla_profesional_tym_word(documento_final_word, d_ren_f, ['#', 'Deportista', 'Tiempo Total' if c_key_f=='Completos' else 'CV', 'Natación', 'Bicicleta', 'Trote'])
        h_analisis_f = documento_final_word.add_paragraph('Análisis del Desempeño:'); aplicar_formato_tym_word(h_analisis_f, 13, True)
        for _, fila_f_loop in d_ren_f.iterrows():
            p_n_f = documento_final_word.add_paragraph(f"{fila_f_loop['#']}. {fila_f_loop['Deportista']}"); aplicar_formato_tym_word(p_n_f, 11, True)
            documento_final_word.add_paragraph(generar_comentario(fila_f_loop, c_key_f, fila_f_loop['#']))

    # BLOQUE 3: TOP 15
    for tit_s_f, ico_f, col_m_f, col_t_f in [('TIEMPO GENERAL', '🥇', 'T_Mins', 'Tiempo Total'), ('NATACIÓN', '🏊‍♂️', 'N_Mins', 'Natación'), ('CICLISMO', '🚴', 'B_Mins', 'Bicicleta'), ('TROTE', '🏃‍♂️', 'R_Mins', 'Trote')]:
        documento_final_word.add_page_break()
        h_15_f = documento_final_word.add_heading(f'{ico_f} TOP 15 {tit_s_f}', level=1); aplicar_formato_tym_word(h_15_f, 15, True); documento_final_word.add_paragraph()
        d_15_f = df_semanal_datos[df_semanal_datos[col_m_f] > 0].sort_values(col_m_f, ascending=False).head(15).copy(); d_15_f['#'] = range(1, len(d_15_f) + 1)
        crear_tabla_profesional_tym_word(documento_final_word, d_15_f, ['#', 'Deportista', col_t_f, 'Natación', 'Bicicleta', 'Trote'] if tit_s_f == 'TIEMPO GENERAL' else ['#', 'Deportista', col_t_f, 'Tiempo Total'])
        h_podio_f = documento_final_word.add_paragraph('Análisis del Podio:'); aplicar_formato_tym_word(h_podio_f, 13, True)
        for _, f_p_f in d_15_f.head(3).iterrows():
            p_at_f = documento_final_word.add_paragraph(f"{'🥇' if f_p_f['#']==1 else '🥈' if f_p_f['#']==2 else '🥉'} {f_p_f['Deportista']}"); aplicar_formato_tym_word(p_at_f, 11, True)
            documento_final_word.add_paragraph(generar_comentario(f_p_f, 'General' if tit_s_f == 'TIEMPO GENERAL' else col_t_f, f_p_f['#']))

    # BLOQUE 4: OCR
    documento_final_word.add_page_break()
    h_d_w_f = documento_final_word.add_heading('📏 PODIO DISTANCIA TOTAL', level=1); aplicar_formato_tym_word(h_d_w_f, 15, True); documento_final_word.add_paragraph()
    for idx_d_f, item_d_f in enumerate(podio_d_lista): documento_final_word.add_paragraph(f"{idx_d_f+1}. {item_d_f['nombre']} ({item_d_f['valor']} km)")
    documento_final_word.add_paragraph(); h_l_w_f = documento_final_word.add_heading('⏱️ PODIO ACTIVIDAD MÁS LARGA', level=1); aplicar_formato_tym_word(h_l_w_f, 15, True); documento_final_word.add_paragraph()
    for idx_l_f, item_l_f in enumerate(podio_l_lista): documento_final_word.add_paragraph(f"{idx_l_f+1}. {item_l_f['nombre']} ({item_l_f['valor']})")
    
    stream_salida_word = io.BytesIO(); documento_final_word.save(stream_salida_word); stream_salida_word.seek(0); return stream_salida_word

# =============================================================================
# SECCIÓN INYECTADA: MOTOR NARRATIVO INDIVIDUAL (ESTILO COLAB V17.0 - FIX NUMÉRICO)
# =============================================================================
def generar_reporte_narrativo_individual(atleta_nom, df_actual, dict_historicos, sem_n):
    """Genera reporte personal con insights narrativos basados en Colab V17.0."""
    doc_p = Document()
    match_key = clean_string(atleta_nom)
    
    # Header Institucional
    p_h = doc_p.add_heading(f'Análisis de Rendimiento Personal: {atleta_nom}', 0)
    aplicar_formato_tym_word(p_h, 18, True, True)
    doc_p.add_paragraph(f"Semana de Entrenamiento: {sem_n}").alignment = 1
    
    # Localizar datos del atleta en el procesamiento actual
    df_actual['MatchKey'] = df_actual['Deportista'].apply(clean_string)
    row_act = df_actual[df_actual['MatchKey'] == match_key]
    if row_act.empty: return None
    row_act = row_act.iloc[0]

    def get_benchmarks(hoja_key, match_k):
        df_h = dict_historicos.get(hoja_key)
        if df_h is None: return 0, 0
        df_h['MatchKey'] = df_h.iloc[:, 0].astype(str).apply(clean_string)
        
        # Identificar columnas de semanas (ej: "Sem 01", "Sem 02")
        cols_sem = [c for c in df_h.columns if str(c).startswith("Sem ")]
        
        # Promedio del Equipo (Mins actuales)
        prefijo = "N" if "Natación" in hoja_key else "B" if "Ciclismo" in hoja_key else "R" if "Trote" in hoja_key else "T"
        avg_equipo = df_actual[f"{prefijo}_Mins"].mean()
        
        # --- FIX DE SEGURIDAD PARA PROMEDIO HISTÓRICO ---
        r_atleta = df_h[df_h['MatchKey'] == match_k]
        if not r_atleta.empty:
            # Extraemos la fila de semanas y forzamos a que todo sea número (lo no numérico será NaN)
            datos_semanas = pd.to_numeric(r_atleta[cols_sem].iloc[0], errors='coerce')
            # Calculamos el promedio ignorando los NaN
            avg_hist = datos_semanas.mean() * 1440 if not pd.isna(datos_semanas.mean()) else 0
        else:
            avg_hist = 0
            
        return avg_equipo, avg_hist

    # Construcción de comparativas por disciplina
    for tit_m, hoja_m, col_m in [("TIEMPO TOTAL", "Tiempo Total", "T_Mins"), ("NATACIÓN", "Natación", "N_Mins"), ("CICLISMO", "Ciclismo", "B_Mins"), ("TROTE", "Trote", "R_Mins")]:
        doc_p.add_heading(tit_m, level=1)
        val_act = row_act[col_m]
        bench_eq, bench_hi = get_benchmarks(hoja_m, match_key)
        
        p_m = doc_p.add_paragraph()
        p_m.add_run(f"Volumen actual: {to_hhmmss_display(val_act)}\n").bold = True
        
        diff_eq = val_act - bench_eq
        txt_eq = f"Rendiste {to_hhmmss_display(abs(diff_eq))} {'MÁS' if diff_eq > 0 else 'MENOS'} que el promedio del equipo."
        run_eq = p_m.add_run(txt_eq)
        run_eq.font.color.rgb = RGBColor(0, 100, 0) if diff_eq >= 0 else RGBColor(180, 0, 0)
        
        diff_hi = val_act - bench_hi
        p_m.add_run(f"\nVs Tu Media Histórica: {to_hhmmss_display(abs(diff_hi))} {'MÁS' if diff_hi > 0 else 'MENOS'}.")

    # Gráfico de barras personal
    fig_p, ax_p = plt.subplots(figsize=(5,3))
    ax_p.bar(['Nat', 'Bici', 'Tro'], [row_act['N_Mins'], row_act['B_Mins'], row_act['R_Mins']], color=['#1E90FF', '#32CD32', '#FF4500'])
    ax_p.set_title("Tu Distribución de Carga (Minutos)")
    buf_p = io.BytesIO(); plt.savefig(buf_p, format='png', bbox_inches='tight'); plt.close(fig_p)
    doc_p.add_paragraph().add_run().add_picture(buf_p, width=Inches(3.5))
    
    doc_p.add_paragraph("─" * 50)
    doc_p.add_paragraph("Generado por Agente TYM 2026").alignment = 2
    b_out = io.BytesIO(); doc_p.save(b_out); b_out.seek(0); return b_out

# *****************************************************************************
# --- 7. INTERFAZ DE USUARIO (STREMLIT) - VERSIÓN PERSISTENTE V2.2.31 ---
# *****************************************************************************

# 1. PANEL DE CONTROL (SIDEBAR)
with st.sidebar:
    st.image("https://raw.githubusercontent.com/dagoperez/reportes-tym/main/logo_tym.png", width=150)
    st.title("Configuración")
    
    st.header("📂 Cargas Maestras")
    cargador_maestro_excel = st.file_uploader("📂 Sube el archivo histórico", type=["xlsx"])
    cargador_semana_excel = st.file_uploader("🏃 Sube Excel de la Semana (Real)", type=["xlsx"])
    num_semana_procesar = st.text_input("Número de Semana (Ej: 08):", "08")
    
    st.divider()
    st.subheader("2️⃣ Planificación (Metas)")
    file_plan = st.file_uploader("👤 Subir Plan Individual (Excel)", type=['xlsx'])
    
    with st.expander("🌍 Metas Globales (Manual)"):
        p_n_h = st.number_input("Natación (Hrs)", value=3.0, step=0.5)
        p_n_s = st.number_input("Natación (Ses)", value=3, step=1)
        p_b_h = st.number_input("Ciclismo (Hrs)", value=4.0, step=0.5)
        p_b_s = st.number_input("Ciclismo (Ses)", value=3, step=1)
        p_t_h = st.number_input("Trote (Hrs)", value=1.5, step=0.5)
        p_t_s = st.number_input("Trote (Ses)", value=2, step=1)

    dict_plan_global = {
        'Natacion_Hrs_Plan': p_n_h, 'Natacion_Ses_Plan': p_n_s,
        'Ciclismo_Hrs_Plan': p_b_h, 'Ciclismo_Ses_Plan': p_b_s,
        'Trote_Hrs_Plan': p_t_h, 'Trote_Ses_Plan': p_t_s
    }

    df_plan_indiv = None
    if file_plan:
        df_temp = pd.read_excel(file_plan)
        pos_id_p = ['Deportista', 'Nombre', 'Atleta']
        c_id_p = next((c for c in df_temp.columns if c in pos_id_p), None)
        if c_id_p:
            df_temp.rename(columns={c_id_p: 'Deportista'}, inplace=True)
            df_plan_indiv = df_temp

    ejecutar = st.button("🚀 PROCESAR JORNADA", use_container_width=True)

# 2. PROCESAMIENTO
if ejecutar:
    if cargador_maestro_excel and cargador_semana_excel:
        try:
            df_semana_raw = pd.read_excel(cargador_semana_excel)
            # Ejecutar Motor 3B
            df_resultados = calcular_metricas_cumplimiento(df_semana_raw, df_plan_indiv, dict_plan_global)
            st.session_state['df_resultados'] = df_resultados
            st.session_state['procesado_ok'] = True
        except Exception as e:
            st.error(f"Error en el procesamiento: {e}")
    else:
        st.error("Faltan archivos obligatorios.")

# 3. VISUALIZACIÓN
if st.session_state.get('procesado_ok'):
    df_res = st.session_state['df_resultados']

    st.header(f"📊 Análisis de Adherencia - Semana {num_semana_procesar}")
    
    # KPIs Visuales con manejo de errores
    m1, m2, m3 = st.columns(3)
    if 'TPI_Global' in df_res.columns:
        m1.metric("TPI Promedio Club", f"{df_res['TPI_Global'].mean():.1f}%")
    if 'Estado_Cumplimiento' in df_res.columns:
        m2.metric("Atletas en Óptimo", len(df_res[df_res['Estado_Cumplimiento'] == 'Óptimo']))
        m3.metric("Atletas en Riesgo", len(df_res[df_res['Estado_Cumplimiento'] == 'Riesgo']))

    st.subheader("📋 Tabla de Cumplimiento Detallado (TPI)")
    
    # --- FILTRO DINÁMICO DE COLUMNAS PARA EVITAR KEYERROR ---
    cols_deseadas = ['Deportista', 'TPI_Global', 'Estado_Cumplimiento', 'Nota_Coach', 'VCI_Global', 'SEI_Global']
    cols_presentes = [c for c in cols_deseadas if c in df_res.columns]
    cols_formato = [c for c in ['TPI_Global', 'VCI_Global', 'SEI_Global'] if c in df_res.columns]

    if cols_presentes:
        st.dataframe(
            df_res[cols_presentes].sort_values(cols_presentes[1] if len(cols_presentes)>1 else cols_presentes[0], ascending=False).style.format("{:.1f}%", subset=cols_formato), 
            use_container_width=True
        )
    else:
        st.warning("No se pudieron generar las columnas de cumplimiento. Revisa los nombres de las columnas en tu Excel semanal.")

    st.divider()
    st.subheader("📥 Descargar Entregables")
    c1, c2 = st.columns(2)
    with c1:
        st.download_button("📄 REPORTE WORD GRUPAL", generar_reporte_word_tym_completo(df_res, num_semana_procesar, "N/A", "N/A"), f"Reporte_Sem_{num_semana_procesar}.docx", use_container_width=True)
    with c2:
        st.download_button("📊 EXCEL HISTÓRICO ACTUALIZADO", crear_excel_actualizado(cargador_maestro_excel, df_res, num_semana_procesar), f"Historico_S{num_semana_procesar}.xlsx", use_container_width=True)

    st.divider()
    st.subheader("📦 Reportes Individuales (ZIP)")
    
    # Preparación de histórico
    try:
        dict_h_ref = {
            "Tiempo Total": pd.read_excel(cargador_maestro_excel, sheet_name="Tiempo Total", dtype=object),
            "Natación": pd.read_excel(cargador_maestro_excel, sheet_name="Natación", dtype=object),
            "Ciclismo": pd.read_excel(cargador_maestro_excel, sheet_name="Ciclismo", dtype=object),
            "Trote": pd.read_excel(cargador_maestro_excel, sheet_name="Trote", dtype=object)
        }
        
        atletas_list = df_res['Deportista'].tolist() if 'Deportista' in df_res.columns else []
        seleccion = st.multiselect("Seleccionar atletas:", atletas_list, default=atletas_list)
        
        if seleccion and st.button("⚙️ GENERAR ARCHIVOS ZIP"):
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                for atleta in seleccion:
                    r_indiv = generar_reporte_narrativo_individual(atleta, df_res, dict_h_ref, num_semana_procesar)
                    if r_indiv:
                        zf.writestr(f"Reporte_{clean_string(atleta)}.docx", r_indiv.getvalue())
            
            st.download_button("⬇️ DESCARGAR ZIP", zip_buffer.getvalue(), f"Individuales_S{num_semana_procesar}.zip")
    except Exception as e:
        st.error(f"Error al preparar el ZIP: {e}. Verifica que el Excel histórico tenga las pestañas: Tiempo Total, Natación, Ciclismo, Trote.")
